# -*- coding: utf-8 -*-
"""Minimal Markdown -> .docx converter for the manuscript (headings, bold/italic,
tables, bullet lists, paragraphs, images, and red callouts).
Usage: python md_to_docx.py <in.md> <out.docx>

Extensions over plain Markdown:
  * Images:        ![Caption text](relative/path.png)   -> centred picture + italic caption
  * Red block:     :::red ... :::                        -> paragraphs rendered in red
  * Inline red:    {{red text}}                          -> red run inside any paragraph
"""
import sys, re, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RED = RGBColor(0xC0, 0x00, 0x00)

def add_runs(par, text, color=None):
    # handle **bold**, *italic*, `code`, {{red}}; strip markdown tokens
    tokens = re.split(r'(\*\*.+?\*\*|\*.+?\*|`.+?`|\{\{.+?\}\})', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('{{') and tok.endswith('}}'):
            r = par.add_run(tok[2:-2]); r.font.color.rgb = RED
        elif tok.startswith('*') and tok.endswith('*'):
            r = par.add_run(tok[1:-1]); r.italic = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = par.add_run(tok[1:-1]); r.font.name = 'Consolas'
        else:
            r = par.add_run(tok)
        if color is not None and r.font.color.rgb is None:
            r.font.color.rgb = color

def main(src, dst):
    base = os.path.dirname(os.path.abspath(src))
    lines = open(src, encoding='utf-8').read().splitlines()
    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    i = 0
    while i < len(lines):
        ln = lines[i].rstrip()
        if not ln.strip():
            i += 1; continue
        if ln.strip() == '---':
            i += 1; continue
        # image:  ![Caption](path)
        mimg = re.match(r'^!\[(.*?)\]\((.+?)\)\s*$', ln.strip())
        if mimg:
            caption, path = mimg.group(1), mimg.group(2).strip()
            ipath = path if os.path.isabs(path) else os.path.join(base, path)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                p.add_run().add_picture(ipath, width=Inches(5.6))
            except Exception as e:
                add_runs(p, f'[image not found: {path}]')
            if caption:
                cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(caption); cr.italic = True; cr.font.size = Pt(9.5)
                cr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1; continue
        # red callout block:  :::red ... :::
        if ln.strip().lower() == ':::red':
            i += 1
            while i < len(lines) and lines[i].strip() != ':::':
                btxt = lines[i].rstrip()
                if btxt.strip():
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    add_runs(p, btxt, color=RED)
                i += 1
            i += 1  # skip closing :::
            continue
        # table block
        if ln.lstrip().startswith('|'):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith('|'):
                block.append(lines[i].strip()); i += 1
            rows = [[c.strip() for c in r.strip('|').split('|')] for r in block]
            rows = [r for r in rows if not all(set(c) <= set('-: ') for c in r)]  # drop separator
            if rows:
                ncol = max(len(r) for r in rows)
                tbl = doc.add_table(rows=0, cols=ncol); tbl.style = 'Light Grid Accent 1'
                for ridx, r in enumerate(rows):
                    cells = tbl.add_row().cells
                    for cidx in range(ncol):
                        txt = r[cidx] if cidx < len(r) else ''
                        cells[cidx].text = ''
                        p = cells[cidx].paragraphs[0]
                        add_runs(p, txt)
                        if ridx == 0:
                            for rr in p.runs: rr.bold = True
            continue
        # headings
        m = re.match(r'^(#{1,6})\s+(.*)$', ln)
        if m:
            level = len(m.group(1)); txt = m.group(2).strip()
            txt = re.sub(r'\*\*(.+?)\*\*', r'\1', txt)
            if level == 1:
                h = doc.add_heading('', level=0); add_runs(h, txt)
            else:
                h = doc.add_heading('', level=min(level-1,4)); add_runs(h, txt)
            i += 1; continue
        # bullet
        if re.match(r'^\s*[-*]\s+', ln):
            txt = re.sub(r'^\s*[-*]\s+', '', ln)
            p = doc.add_paragraph(style='List Bullet'); add_runs(p, txt)
            i += 1; continue
        # blockquote
        if ln.lstrip().startswith('>'):
            txt = re.sub(r'^\s*>\s?', '', ln)
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(18); add_runs(p, txt)
            i += 1; continue
        # normal paragraph (gather until blank)
        para = [ln]; i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].lstrip().startswith(('|','#','-','*','>','!',':::')) and lines[i].strip() != '---':
            para.append(lines[i].rstrip()); i += 1
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, ' '.join(para))
    doc.save(dst)
    print('Saved', dst)

if __name__ == '__main__':
    main(sys.argv[1], sys.argv[2])
